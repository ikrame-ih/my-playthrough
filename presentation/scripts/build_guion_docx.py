# Genera guion-defensa-8min.docx en presentation/docs/
"""Guion con tono alumna: frases naturales, fáciles de estudiar. 8 min exposición + ~2 min tribunal."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "guion-defensa-8min.docx"


def set_para_spacing(p, before=0, after=0):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    set_para_spacing(p, before=10, after=3)
    return p


def add_para(doc: Document, text: str, bold: bool = False, color: RGBColor = None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    if color:
        run.font.color.rgb = color
    set_para_spacing(p, before=1, after=1)
    return p


def add_time_badge(doc: Document, label: str):
    p = doc.add_paragraph()
    run = p.add_run("Tiempo: " + label)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)  # slate-500 — legible
    set_para_spacing(p, before=2, after=2)


def add_note(doc: Document, text: str):
    """Nota de dirección de escena (gris oscuro legible, entre corchetes)."""
    p = doc.add_paragraph()
    run = p.add_run(f"[{text}]")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)  # slate-600 — gris oscuro
    set_para_spacing(p, before=1, after=1)
    return p


def add_speech(doc: Document, body: str):
    """Texto del discurso en cursiva sobre fondo blanco de Word."""
    p = doc.add_paragraph()
    run = p.add_run(body)
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)  # slate-900 — legible sobre blanco
    set_para_spacing(p, before=2, after=2)
    return p


def add_divider(doc: Document):
    p = doc.add_paragraph("─" * 60)
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(0x33, 0x3A, 0x4A)
    set_para_spacing(p, before=6, after=6)


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    # ── Título ──────────────────────────────────────────────────────────────
    t = doc.add_heading("Guion de defensa · MyPlaythrough", 0)
    t.runs[0].font.size = Pt(20)
    set_para_spacing(t, before=0, after=6)

    add_para(
        doc,
        "Ikrame Ibn Hayoun · 2.º DAW · CESUR Málaga Este · Proyecto integrado PERN · Defensa 5 mayo 2026.",
        color=RGBColor(0x6B, 0x72, 0x80),
    )

    add_divider(doc)

    # ── Cómo usarlo ─────────────────────────────────────────────────────────
    add_heading(doc, "Cómo usarlo", 1)
    add_para(
        doc,
        "Las frases están escritas como las diría una alumna, no como un manual. "
        "No tienes que memorizar cada palabra: lee en voz alta dos o tres veces y "
        "quédate con el orden y las dos ideas clave de cada slide. "
        "Si te bloqueas, parafrasea; el tribunal evalúa que entiendes lo que has hecho, "
        "no que hayas memorizado un texto.",
    )

    add_heading(doc, "Cronograma — suman exactamente 8:00", 2)
    for line in [
        "0:00 → 0:40  —  Slide 1 · Portada",
        "0:40 → 2:05  —  Slide 2 · Qué es (+ comparativa con otras apps)",
        "2:05 → 3:10  —  Slide 3 · Arquitectura PERN (+ por qué este stack)",
        "3:10 → 4:25  —  Slide 4 · Modelo de datos",
        "4:25 → 5:05  —  Slide 5 · Funciones principales",
        "5:05 → 7:40  —  Slide 6 · Demo en vivo",
        "7:40 → 8:00  —  Slide 7 · Cierre",
        "8:00 →  ~     —  Preguntas del tribunal (~2 min, no cuenta en tus 8)",
    ]:
        add_para(doc, "   " + line)

    add_para(
        doc,
        "IMPORTANTE: Cuando la profesora levante la mano en GoToMeeting (minuto 6) "
        "ya sabes que te quedan 2 minutos. Termina la demo de forma natural y pasa al cierre.",
        bold=True,
        color=RGBColor(0xFB, 0xBF, 0x24),
    )

    add_divider(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # SLIDE 1 — PORTADA
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "SLIDE 1  ·  Portada", 1)
    add_time_badge(doc, "0:00 → 0:40  (40 segundos)")

    add_note(
        doc,
        "La slide muestra el nombre del proyecto, tus datos académicos, la tutora y el centro. "
        "No la leas. Mira a cámara y habla directamente.",
    )

    add_speech(
        doc,
        "«Buenas tardes. Soy Ikrame, de segundo de DAW en CESUR Málaga Este. "
        "Mi proyecto se llama MyPlaythrough: es una aplicación web para gestionar "
        "tu colección personal de videojuegos y tener una parte de comunidad. "
        "Está construida con el stack PERN —PostgreSQL, Express, React y Node—. "
        "Os cuento de qué va, cómo lo he montado técnicamente, cómo está la base de datos, "
        "qué funciones tiene, y para terminar os la enseño en el navegador.»",
    )

    add_divider(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # SLIDE 2 — QUÉ ES / COMPARATIVA
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "SLIDE 2  ·  Qué es MyPlaythrough", 1)
    add_time_badge(doc, "0:40 → 2:05  (1 min 25 seg)")

    add_note(
        doc,
        "La slide tiene dos tarjetas: 'Colección' (estado, notas, horas, portadas) "
        "y 'Comunidad' (perfiles, seguir, recomendar, LFG). También hay una línea "
        "sobre Steam y RAWG, y la paleta de colores. Úsala de apoyo visual, no la leas.",
    )

    add_para(doc, "Punto 1 — El problema y la comparativa:", bold=True)
    add_speech(
        doc,
        "«La idea surge de un problema bastante común si juegas: tienes un montón de juegos "
        "empezados, otros que llevas años sin tocar, algunos completados… y si no los apuntas "
        "en algún sitio los olvidas. "
        "Hay apps que ya hacen algo parecido: Backloggd es la más conocida, pero carga lenta, "
        "tiene bastante ruido visual y para las cosas simples resulta más de lo que necesitas. "
        "HowLongToBeat es buena para saber cuántas horas tiene un juego, pero no es para "
        "llevar TU biblioteca ni apuntar tus notas personales. "
        "La biblioteca de Steam solo sirve si tienes el juego en Steam; "
        "si juegas en Nintendo, PlayStation o en cualquier otra plataforma, no aparece ahí.»",
    )

    add_para(doc, "Punto 2 — La propuesta de MyPlaythrough:", bold=True)
    add_speech(
        doc,
        "«MyPlaythrough lo que hace es ser más limpio: añades cualquier juego de cualquier plataforma, "
        "le pones tu estado —jugando, completado, pendiente, abandonado—, tus horas, una nota personal, "
        "y la portada la busca solo con Steam y RAWG. Sin publicidad, sin ruido. "
        "Y tiene una parte de comunidad que no abruma: puedes seguir a quien te interese, "
        "recomendar un juego a alguien que ya sigues, y buscar grupo si quieres jugar con alguien.»",
    )

    add_para(doc, "Punto 3 — Diseño (unos segundos):", bold=True)
    add_speech(
        doc,
        "«El fondo oscuro no es solo estético: las portadas de los juegos resaltan mucho más "
        "que sobre fondo blanco, y si estás mirando tu lista un buen rato no cansa la vista. "
        "El color de acción principal es ese verde agua que veis en la diapositiva.»",
    )

    add_divider(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # SLIDE 3 — ARQUITECTURA / POR QUÉ PERN
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "SLIDE 3  ·  Arquitectura · PERN", 1)
    add_time_badge(doc, "2:05 → 3:10  (1 min 5 seg)")

    add_note(
        doc,
        "La slide muestra tres cajas: Cliente (React / Vite / Tailwind), "
        "API (Node / Express / REST) y Base de datos (PostgreSQL), "
        "con flechas animadas entre ellas. Señálalas de izquierda a derecha mientras hablas.",
    )

    add_para(doc, "Punto 1 — Capa de cliente (React + Vite + Tailwind):", bold=True)
    add_speech(
        doc,
        "«En el navegador tengo React con Vite. "
        "React me deja dividir la interfaz en componentes reutilizables "
        "—la barra lateral, una tarjeta de juego, el modal de recomendaciones— "
        "y la pantalla se actualiza solo en la parte que cambia, sin recargar todo. "
        "Vite arranca y compila mucho más rápido que Create React App. "
        "Tailwind me da las clases de estilo directamente en el HTML "
        "para no tener que saltar entre archivos CSS y JSX continuamente.»",
    )

    add_para(doc, "Punto 2 — Por qué no PHP+MySQL (lo que usa mucha gente en DAW):", bold=True)
    add_speech(
        doc,
        "«Podría haberlo hecho con PHP y MySQL, que también hemos visto en el ciclo, "
        "pero elegí JavaScript en los dos lados —cliente y servidor— para no cambiar "
        "de lenguaje cada vez que pasas del front al back. Todo el proyecto es JS/TS, "
        "se despliega fácil con Docker y el ecosistema de paquetes npm es enorme.»",
    )

    add_para(doc, "Punto 3 — Capa de servidor (Node + Express):", bold=True)
    add_speech(
        doc,
        "«Express sobre Node me da una API REST limpia: defino la ruta, el middleware "
        "que comprueba el token, y la lógica. "
        "Las portadas de los juegos pasan también por el servidor como proxy "
        "para que el navegador no tenga problemas pidiendo imágenes de dominios externos. "
        "Y tengo un límite de intentos de login por IP para que no se pueda ir probando "
        "contraseñas sin freno.»",
    )

    add_para(doc, "Punto 4 — Por qué PostgreSQL y no MongoDB:", bold=True)
    add_speech(
        doc,
        "«Usé PostgreSQL porque los datos son relacionales: "
        "una cuenta tiene muchos juegos, cada juego referencia una ficha del catálogo "
        "que comparten varias personas, los seguimientos son un grafo de usuarios… "
        "Eso encaja perfecto en tablas con claves foráneas e integridad referencial. "
        "MongoDB sería más libre pero perdería las garantías que da SQL sobre "
        "que los datos no quedan inconsistentes.»",
    )

    add_divider(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # SLIDE 4 — MODELO DE DATOS
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "SLIDE 4  ·  Modelo de datos", 1)
    add_time_badge(doc, "3:10 → 4:25  (1 min 15 seg)")

    add_note(
        doc,
        "La slide tiene arriba una frase resumen: '1 cuenta → muchas filas en tu lista → "
        "cada una apunta a 1 ficha del catálogo (compartida)'. "
        "Luego tres cajas: usuarios → juegos → catalogo_juegos. "
        "Debajo, cuatro tarjetas de la parte social. "
        "Señala las cajas de izquierda a derecha y la parte de abajo al mencionarla.",
    )

    add_para(doc, "Idea principal — explicar el problema que resuelve el diseño:", bold=True)
    add_speech(
        doc,
        "«Lo que me parece más interesante de la base de datos es la decisión sobre las portadas. "
        "Imagina que diez personas añaden The Witcher 3 a su biblioteca. "
        "Si guardara el título y la portada en cada fila de cada usuario, "
        "estaría duplicando la misma información diez veces. "
        "Lo que hice es separarlo: existe una tabla llamada catalogo_juegos "
        "donde se guarda la ficha del juego —nombre, portada, el ID de RAWG o Steam—, "
        "y cada usuario tiene una fila en 'juegos' que apunta a esa ficha compartida. "
        "Así el título y la imagen están una sola vez, y lo que es personal "
        "—tu estado, tus horas, tu nota— está en tu fila.»",
    )

    add_para(doc, "Las tres tablas del núcleo:", bold=True)
    add_speech(
        doc,
        "«La tabla 'usuarios' guarda todo lo de la cuenta: el email, la contraseña hasheada, "
        "el nombre público, el avatar, el rol —usuario o administrador—. "
        "De ahí cuelga todo lo que es tuyo. "
        "La tabla 'juegos' es tu biblioteca: una fila por juego que hayas añadido, "
        "con tu estado, horas, nota personal y un enlace a la ficha del catálogo. "
        "Y 'catalogo_juegos' es la ficha compartida: título, portada, "
        "y los IDs de RAWG y Steam para no buscarla dos veces.»",
    )

    add_para(doc, "La parte social:", bold=True)
    add_speech(
        doc,
        "«Abajo en la slide hay cuatro tarjetas más: comentarios con votos sobre la ficha pública "
        "de un juego, seguimientos entre usuarios —que es básicamente un grafo—, "
        "recomendaciones que llegan a tu bandeja, "
        "y publicaciones de buscar grupo vinculadas a un juego de tu biblioteca. "
        "Todo enlazado con claves foráneas, y si borras un usuario, "
        "lo suyo se borra en cascada sin dejar datos huérfanos.»",
    )

    add_divider(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # SLIDE 5 — FUNCIONES
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "SLIDE 5  ·  Funciones principales", 1)
    add_time_badge(doc, "4:25 → 5:05  (40 segundos)")

    add_note(
        doc,
        "Hay seis tarjetas: Entrar en la app, Tu colección, Perfiles y comunidad, "
        "Recomendaciones, Buscar grupo (LFG) y Administración. "
        "No las leas una a una: pasa rápido y deja que la demo sea quien lo muestre.",
    )

    add_speech(
        doc,
        "«Esta slide es un resumen rápido de las seis áreas principales. "
        "La autenticación tiene validación de contraseña tanto en el servidor como en el formulario. "
        "La colección permite buscar portadas, vista cuadrícula o lista, y al guardar "
        "te avisa con un mensaje visible cuando vuelves a la pantalla principal. "
        "La parte social tiene seguir usuarios, recomendaciones con bandeja y campana, "
        "y buscar grupo si quieres quedar para jugar con alguien. "
        "Y hay un panel de administración para moderar usuarios y contenidos. "
        "Ahora os lo enseño en directo.»",
    )

    add_divider(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # SLIDE 6 — DEMO
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "SLIDE 6  ·  Demo en vivo", 1)
    add_time_badge(doc, "5:05 → 7:40  (2 min 35 seg)")

    add_note(
        doc,
        "ANTES DE EMPEZAR: Tener el navegador preparado con la pestaña de la app abierta, "
        "sesión iniciada como Demo Jurado (demo@myplaythrough.local / Presentacion2026!) "
        "o con otra cuenta con datos. Tener la colección visible. "
        "Si algo falla, sigue hablando con calma y di 'voy a ir a…' sin pánico.",
    )

    add_para(doc, "Bloque 1 — Colección (1 minuto aprox.):", bold=True)
    add_speech(
        doc,
        "«Pasamos a la aplicación en el navegador. "
        "Estoy en la pantalla de inicio, que es mi colección. "
        "Aquí se ven los juegos que tengo añadidos con su portada, el estado y la puntuación. "
        "[Mostrar vista cuadrícula y luego lista si hay tiempo.] "
        "Voy a añadir uno nuevo: le doy a 'Añadir juego', busco un título cualquiera, "
        "[escribir algo en el buscador de portadas, elegir una], "
        "le pongo el estado y guardo. "
        "Al volver a la colección aparece el aviso de confirmación que os comenté; "
        "así sabes que se guardó bien sin tener que buscar la fila.»",
    )

    add_para(doc, "Bloque 2 — Comunidad y social (1 minuto aprox.):", bold=True)
    add_speech(
        doc,
        "«Ahora voy a la sección de Comunidad. "
        "[Navegar a /community.] "
        "Aquí aparecen los miembros registrados, estadísticas globales del catálogo y actividad "
        "de las personas que sigo. "
        "Voy a abrir el perfil de un usuario "
        "[clic en cualquier tarjeta de usuario], "
        "se ve su colección en modo solo lectura. "
        "Y si le doy a recomendar, se abre un modal donde selecciono "
        "qué juego de mi biblioteca mando. "
        "El receptor le llega en la campana de arriba "
        "[señalar el icono de campana en el header], "
        "aquí se ve el contador de no leídas.»",
    )

    add_para(doc, "Bloque 3 — Opcional si queda tiempo (30 seg):", bold=True)
    add_speech(
        doc,
        "«Si tengo un segundo más, os enseño el panel de admin: "
        "[ir a /admin si la cuenta tiene rol admin] "
        "desde aquí se pueden gestionar usuarios, fichas y publicaciones de buscar grupo. "
        "El borrado de cuenta pide escribir el nombre exacto del usuario para confirmar, "
        "para que no sea un clic accidental.»",
    )

    add_para(doc, "Cerrar la demo:", bold=True)
    add_speech(
        doc,
        "«Todo el código, las instrucciones para arrancarlo y la cuenta demo "
        "están documentados en el README del repositorio. Volvemos a las diapositivas.»",
    )

    add_divider(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # SLIDE 7 — CIERRE
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "SLIDE 7  ·  Cierre", 1)
    add_time_badge(doc, "7:40 → 8:00  (20 segundos)")

    add_note(
        doc,
        "La slide muestra '¿Preguntas?', la URL del repositorio y la carpeta docs/. "
        "Breve, directo y con calma.",
    )

    add_speech(
        doc,
        "«Y hasta aquí la presentación. "
        "El código, la documentación y el plan de pruebas están en el repositorio "
        "que aparece en pantalla. "
        "Muchas gracias por la atención; cuando queráis, estoy lista para las preguntas.»",
    )

    add_divider(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # PREGUNTAS DEL TRIBUNAL
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "Preguntas del tribunal  (~2 min, fuera de los 8)", 1)

    add_para(
        doc,
        "Técnica: escucha hasta el final, no interrumpas. Si no entiendes del todo: "
        "'¿Se refiere a…?' en una frase. Responde primero simple, luego técnico si hace falta. "
        "Si no sabes algo, lo dices con naturalidad: 'No lo he implementado así, pero lo haría de esta forma…'",
    )

    qa = [
        (
            "¿Por qué PostgreSQL y no MongoDB o MySQL?",
            "«Mis datos son relacionales: usuarios que tienen juegos, juegos que referencian el catálogo, "
            "seguimientos entre usuarios… Eso encaja perfecto en tablas con claves foráneas. "
            "MongoDB sería más flexible pero perdería la integridad referencial que me da SQL. "
            "Y PostgreSQL en concreto lo elegí porque soporta CASCADE en los borrados, "
            "índices únicos con expresión —para que los nombres de usuario sean únicos sin distinguir "
            "mayúsculas—, y porque lo hemos trabajado en clase.»",
        ),
        (
            "¿Cómo funciona la autenticación / el login?",
            "«Al registrarte, la contraseña se guarda como un hash bcrypt; nunca en claro. "
            "Al hacer login correcto, el servidor genera un JWT —un token firmado con un secreto— "
            "y el navegador lo guarda y lo manda en cada petición siguiente en la cabecera Authorization. "
            "Si el token no es válido o ha expirado, la API devuelve 401 y el cliente te redirige al login. "
            "Los roles se comprueban en base de datos en cada petición de admin, "
            "no solo en lo que diga el token.»",
        ),
        (
            "¿Qué pasa si borras un usuario?",
            "«Usé DELETE CASCADE en las claves foráneas. "
            "Cuando borro un usuario, PostgreSQL borra automáticamente sus juegos, "
            "sus comentarios, sus seguimientos y sus recomendaciones. "
            "Así no quedan datos huérfanos en la base. "
            "En el panel de admin además hay un modal que te pide escribir el nombre "
            "exacto del usuario para confirmar, para evitar borrados accidentales.»",
        ),
        (
            "¿De dónde salen las portadas de los juegos?",
            "«De dos APIs públicas: Steam y RAWG. "
            "Las pido desde el servidor y las sirvo como proxy al cliente, "
            "porque si el navegador las pide directamente a esos dominios a veces falla "
            "por restricciones CORS o de hotlinking. "
            "Además, guardo el ID de RAWG o Steam en catalogo_juegos "
            "para que si varios usuarios añaden el mismo juego, "
            "la portada solo se busca una vez.»",
        ),
        (
            "¿Qué medidas de seguridad has tomado?",
            "«Varias: consultas parametrizadas en PostgreSQL para evitar inyección SQL; "
            "contraseñas hasheadas con bcrypt; JWT para autenticación; "
            "el rol de admin se comprueba en base de datos en cada petición, no solo en el front; "
            "CORS configurado para que solo el origen del frontend pueda llamar a la API; "
            "rate limiting en login y registro para que no se pueda ir probando contraseñas sin freno; "
            "y en producción el endpoint de diagnóstico /api/test-db devuelve 404.»",
        ),
        (
            "¿Qué fue lo más difícil?",
            "«Mantener consistencia entre las tres capas: que lo que pintaba React "
            "fuera exactamente lo que devolvía la API y lo que había en la base. "
            "Especialmente en permisos y en la parte social, "
            "donde una acción —como recomendar o seguir— afecta a varios usuarios "
            "y hay que manejar bien los estados en el front sin que se desfasen.»",
        ),
        (
            "¿Qué mejorarías o añadirías?",
            "«Con más tiempo metería más pruebas automáticas —ahora tengo un plan de pruebas manual "
            "en docs/pruebas.md pero pocos tests automatizados—. "
            "También me gustaría tener despliegue real en un servidor público, "
            "y quizá notificaciones en tiempo real con WebSockets en vez de polling. "
            "Y una función de estadísticas personales más detallada: "
            "tu género favorito, plataforma más usada, ese tipo de cosas.»",
        ),
        (
            "¿Con qué herramienta has hecho las diapositivas?",
            "«Son una mini aplicación aparte del cliente del proyecto, "
            "hecha con React, Vite y TypeScript. "
            "El estilo es Tailwind con los mismos colores que la app. "
            "Las animaciones de entrada y salida de los elementos son con Framer Motion: "
            "un fadeUp con un poco de blur y movimiento. "
            "Para el PDF hay una copia estática separada en HTML plano sin Framer, "
            "porque si imprimes el árbol animado el PDF sale corrupto o vacío. "
            "Los iconos son SVG propios y las tipografías son Syne y Manrope.»",
        ),
        (
            "¿Por qué React y no Vue o Angular?",
            "«React es el más extendido en el mercado laboral y el que más hemos trabajado. "
            "Vue sería una buena alternativa, quizá más sencillo para empezar, "
            "pero React tiene más ecosistema y más recursos. "
            "Angular me parecía demasiado pesado para este tamaño de proyecto, "
            "tiene mucha configuración inicial para lo que necesitaba.»",
        ),
        (
            "¿Qué es un PERN stack?",
            "«Son las siglas de las cuatro tecnologías principales: "
            "PostgreSQL para la base de datos, Express como framework del servidor, "
            "React para el frontend y Node.js como entorno de ejecución del servidor. "
            "Es uno de los stacks más comunes hoy en desarrollo web con JavaScript puro "
            "en los dos lados —cliente y servidor—.»",
        ),
    ]

    for q, a in qa:
        add_para(doc, q, bold=True)
        add_speech(doc, a)
        doc.add_paragraph()

    add_divider(doc)

    # ═══════════════════════════════════════════════════════════════════════
    # INSTRUCCIONES PARA LA IA DE APOYO (CUSTOM PROMPT)
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "Instrucciones para la IA de apoyo (Custom Prompt)", 1)
    add_para(
        doc,
        "Copia el texto de abajo en la herramienta de IA que usas durante la defensa "
        "(campo 'Instructions' del Custom Prompt). Dale un nombre como 'Defensa MyPlaythrough'.",
        color=RGBColor(0xFB, 0xBF, 0x24),
        bold=True,
    )

    prompt_text = (
        "Eres un asistente de apoyo en tiempo real para la defensa de un Proyecto Integrado de DAW. "
        "Tu única función es ayudar a la alumna a responder preguntas del tribunal sobre su proyecto, "
        "llamado MyPlaythrough. Sé conciso, claro y usa el mismo tono natural de la alumna. "
        "No añadas información que no sea del proyecto. "
        "\n\n"
        "=== DESCRIPCIÓN DEL PROYECTO ===\n"
        "MyPlaythrough es una aplicación web SPA (Single Page Application) para gestionar "
        "la colección personal de videojuegos de un usuario y conectar con una pequeña comunidad. "
        "Es el proyecto integrado final del CFGS DAW de Ikrame Ibn Hayoun "
        "en CESUR Málaga Este, curso 2025/2026, tutora María Jesús Rodríguez Sánchez.\n"
        "\n"
        "=== STACK TECNOLÓGICO (PERN) ===\n"
        "- PostgreSQL: base de datos relacional, 8 tablas, integridad referencial, CASCADE delete.\n"
        "- Express 5 + Node.js: API REST en JavaScript. Sin recargas, JSON puro.\n"
        "- React 18 + Vite: frontend SPA. Tailwind CSS para estilos. React Router DOM para rutas.\n"
        "- Autenticación: JWT firmado con JWT_SECRET, roles user/admin comprobados en BD en cada petición.\n"
        "- Contraseñas: hash bcrypt coste 10 (bcryptjs). Nunca en claro.\n"
        "- CORS: solo el origen configurado en CORS_ORIGIN puede llamar a la API.\n"
        "- Rate limiting: express-rate-limit en /api/auth/login y /api/auth/register.\n"
        "- Portadas: proxy en servidor para Steam y RAWG (evita CORS/hotlinking en el browser).\n"
        "\n"
        "=== BASE DE DATOS — 8 TABLAS ===\n"
        "1. usuarios: id, nombre_usuario (ÚNICO, índice lower+trim), email (único), password_hash, "
        "rol (user/admin), avatar_id, notificaciones_sonido, fecha_registro.\n"
        "2. catalogo_juegos: ficha compartida. id, titulo, url_imagen, rawg_id (UNIQUE), "
        "steam_app_id (UNIQUE). Varios usuarios pueden referenciar la misma fila.\n"
        "3. juegos: biblioteca personal. id, usuario_id (FK→usuarios CASCADE), titulo, estado "
        "(jugando/completado/pendiente/abandonado/wishlist), plataforma, puntuacion, horas_jugadas, "
        "comentario, url_imagen, catalogo_id (FK→catalogo_juegos SET NULL).\n"
        "4. juego_comentarios: hilos de reseñas. id, juego_id (FK CASCADE), usuario_id (FK CASCADE), "
        "parent_id (FK self-ref, respuestas), cuerpo, fecha_creacion.\n"
        "5. juego_comentario_votos: voto por usuario por comentario. PK compuesta (comentario_id, usuario_id), "
        "valor SMALLINT (1 o -1). Solo comentarios de primer nivel.\n"
        "6. usuario_seguimientos: grafo de seguir. (seguidor_id, seguido_id) PK compuesta, ambos FK→usuarios.\n"
        "7. juego_recomendaciones: remitente_id, destinatario_id (FK→usuarios), juego_id (FK→juegos), "
        "mensaje, leida, fecha.\n"
        "8. lfg_publicaciones: buscar grupo. usuario_id (FK), juego_id (FK), modo (online/local/otro), "
        "descripcion, activa, fecha.\n"
        "\n"
        "=== DECISIÓN DE DISEÑO CLAVE ===\n"
        "La tabla catalogo_juegos existe para no duplicar título y portada. Si 50 usuarios añaden "
        "'The Witcher 3', hay 50 filas en 'juegos' (una por usuario con sus datos personales) "
        "pero solo 1 fila en 'catalogo_juegos' con el título e imagen. Cada fila de 'juegos' "
        "referencia esa fila compartida mediante catalogo_id.\n"
        "\n"
        "=== FUNCIONALIDADES PRINCIPALES ===\n"
        "- Registro con validación de contraseña (8+ chars, mayúscula, minúscula, número, símbolo) "
        "en servidor y formulario. Nombre público único (case-insensitive). "
        "Login por email O por nombre_usuario.\n"
        "- Colección: CRUD completo de juegos. Búsqueda de portadas (Steam + RAWG). "
        "Vista cuadrícula o lista compacta. Ordenación por título, estado, nota, reciente. "
        "Aviso flash visible al volver a la colección tras guardar/editar (router location.state).\n"
        "- Comunidad: lista de usuarios, seguir/dejar de seguir, estadísticas globales (AVG, COUNT con SQL), "
        "feed de actividad de seguidos (comentarios + LFG).\n"
        "- Recomendaciones: solo puedes recomendar a alguien a quien sigues. "
        "Bandeja en /recommendations. Campana con contador de no leídas. "
        "Tono de notificación opcional (Web Audio API, toggle en Perfil). "
        "Modal confirma destinatario y título antes de enviar.\n"
        "- LFG (buscar grupo): publicación ligada a un juego de tu biblioteca, modo de juego, descripción.\n"
        "- Comentarios y votos: hilos anidados por juego, voto positivo/negativo en comentarios raíz.\n"
        "- Avatares: 10 robots SVG predefinidos, guardado en usuarios.avatar_id.\n"
        "- Panel admin: gestión de usuarios (borrar con modal de confirmación —escribe el nombre exacto—), "
        "fichas y LFG. Rol comprobado en BD en cada petición admin. /admin redirige al inicio si no es admin.\n"
        "- Tour guiado (react-joyride): se muestra una vez en ese navegador tras el primer registro. "
        "Se puede relanzar desde Perfil. Guardado en localStorage por userId.\n"
        "- Skeletons: placeholders animados mientras cargan los datos.\n"
        "- Accesibilidad: skip link, foco a <main>, labels en controles.\n"
        "\n"
        "=== ARQUITECTURA DEL SERVIDOR ===\n"
        "server/index.js → monta routers por tema:\n"
        "  /api/auth   → auth.routes.js (login, register, GET/PATCH me)\n"
        "  /api/games  → games.routes.js (CRUD colección) + covers.routes.js (proxy portadas)\n"
        "  /api/users  → users.routes.js (perfiles públicos)\n"
        "  /api/community → community.routes.js (stats, actividad)\n"
        "  /api/social → social.routes.js (follows, recomendaciones, LFG)\n"
        "  /api/admin  → admin.routes.js (moderación)\n"
        "Middlewares: authMiddleware (JWT), adminMiddleware (rol en BD), express-rate-limit.\n"
        "\n"
        "=== SEGURIDAD (RESUMEN RÁPIDO) ===\n"
        "- SQL injection: consultas parametrizadas ($1, $2) siempre, nunca concatenación.\n"
        "- Passwords: bcrypt hash (coste 10), nunca plaintext.\n"
        "- JWT: firmado con JWT_SECRET, expira en 7 días.\n"
        "- Admin: rol comprobado en BD en cada petición (no solo token).\n"
        "- CORS: origen exacto configurado en .env.\n"
        "- Rate limiting: 40 intentos/15 min en auth.\n"
        "- Body limit: 50kb max en express.json.\n"
        "- Proxy imágenes: solo dominios permitidos (Steam, RAWG).\n"
        "- /api/test-db devuelve 404 en producción.\n"
        "\n"
        "=== PUESTA EN MARCHA ===\n"
        "Opción A (Docker): docker compose up --build — levanta API en :3000 y Postgres en host :5433.\n"
        "Opción B (Manual): crear BD con docs/sql/schema.sql, configurar server/.env "
        "(DB_*, JWT_SECRET, CORS_ORIGIN, RAWG_API_KEY opcional), npm run dev en server/ y client/.\n"
        "Datos de demo: npm run seed:presentation (desde server/) — crea 7 usuarios "
        "con contraseña Presentacion2026!. Login rápido en pantalla de auth: botón 'Rellenar cuenta demo' "
        "→ usa Demo Jurado (demo@myplaythrough.local).\n"
        "\n"
        "=== DIAPOSITIVAS (orden de la presentación) ===\n"
        "1. Portada — nombre, datos académicos.\n"
        "2. Qué es — colección + comunidad, comparativa con Backloggd/HowLongToBeat/Steam, diseño oscuro.\n"
        "3. Arquitectura PERN — React/Vite/Tailwind → Express/Node → PostgreSQL, flechas animadas.\n"
        "4. Modelo de datos — 3 tablas núcleo + 4 de comunidad, decisión catalogo_juegos separado.\n"
        "5. Funciones — 6 tarjetas (auth, colección, comunidad, recomendaciones, LFG, admin).\n"
        "6. Demo en vivo — colección, búsqueda portada, social, campana, admin.\n"
        "7. Cierre — URL repositorio, docs/.\n"
        "\n"
        "=== PREGUNTAS FRECUENTES RESUMIDAS ===\n"
        "- ¿Por qué PERN y no MERN? Datos relacionales → PostgreSQL > MongoDB.\n"
        "- ¿Por qué no PHP+MySQL? JS en los dos lados, ecosistema npm, Docker fácil.\n"
        "- ¿Por qué React y no Vue/Angular? Más mercado, más usado en clase, componentes reutilizables.\n"
        "- ¿Por qué catalogo_juegos separado? Para no duplicar título+portada por cada usuario.\n"
        "- ¿JWT o sesiones? JWT: sin estado en servidor, fácil de escalar, estándar en APIs REST.\n"
        "- ¿Qué mejorarías? Más tests automáticos, despliegue real, WebSockets para notificaciones, "
        "estadísticas personales avanzadas.\n"
        "- ¿Qué costó más? Consistencia entre las 3 capas en permisos y social.\n"
        "- ¿Las diapositivas? Mini app React+Vite+TypeScript+Tailwind+Framer Motion aparte; "
        "PDF con versión estática sin Framer (PrintDeck).\n"
    )

    add_para(doc, "─── INICIO DEL PROMPT (copia desde aquí hasta FIN DEL PROMPT) ───",
             color=RGBColor(0xFB, 0xBF, 0x24))
    p = doc.add_paragraph()
    run = p.add_run(prompt_text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
    set_para_spacing(p, before=2, after=2)
    add_para(doc, "─── FIN DEL PROMPT ───", color=RGBColor(0xFB, 0xBF, 0x24))

    # Guardar
    try:
        doc.save(OUT)
        print(f"Escrito: {OUT}")
    except PermissionError:
        alt = OUT.with_name("guion-defensa-8min-actualizado.docx")
        doc.save(alt)
        print(f"No se pudo sobrescribir {OUT} (¿abierto en Word?). Guardado como: {alt}")


if __name__ == "__main__":
    main()
