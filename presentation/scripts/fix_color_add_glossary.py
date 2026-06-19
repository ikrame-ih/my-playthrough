# fix_color_add_glossary.py
# Abre el guion existente (preserva cambios manuales), corrige el color del discurso
# (blanco/slate-200 → negro legible) y añade el glosario al final.

from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_LINE_SPACING

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "docs" / "guion-defensa-8min.docx"

# ─── Colores del texto legible sobre fondo blanco ────────────────────────────
COLOR_SPEECH    = RGBColor(0x1E, 0x29, 0x3B)   # slate-900 — discurso (antes era blanco)
COLOR_NOTE      = RGBColor(0x4B, 0x55, 0x63)   # slate-600 — notas de dirección
COLOR_TIME      = RGBColor(0x64, 0x74, 0x8B)   # slate-500 — badges de tiempo
COLOR_HEADING   = RGBColor(0x0F, 0x17, 0x2A)   # casi negro — encabezados
COLOR_YELLOW    = RGBColor(0xB4, 0x53, 0x09)   # ámbar oscuro — avisos en amarillo
COLOR_TERM      = RGBColor(0x1D, 0x4E, 0xD8)   # azul oscuro — término del glosario

# Colores "invisibles" que hay que arreglar (tonos muy claros / blancos)
COLORS_TO_FIX = {
    (0xE2, 0xE8, 0xF0),   # slate-200 — discurso original
    (0xCB, 0xD5, 0xE1),   # slate-300 — prompt de IA
    (0xFF, 0xFF, 0xFF),   # blanco puro
}


def _brightness(r, g, b):
    """Luminancia perceptiva 0-255."""
    return 0.299 * r + 0.587 * g + 0.114 * b


def fix_light_colors(doc: Document):
    """Recorre todos los runs y oscurece colores demasiado claros (luminancia > 180)."""
    fixed = 0
    for para in doc.paragraphs:
        for run in para.runs:
            try:
                rgb = run.font.color.rgb
                if rgb is None:
                    continue
                r, g, b = int(rgb[0]), int(rgb[1]), int(rgb[2])
                # Solo tocar si es un tono muy claro (casi blanco / slate-200 o más claro)
                if _brightness(r, g, b) > 180:
                    run.font.color.rgb = COLOR_SPEECH
                    fixed += 1
            except Exception:
                pass
    print(f"  Runs con color corregido: {fixed}")


# ─── Helpers para añadir contenido ───────────────────────────────────────────

def sp(p, before=0, after=0):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE


def add_h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    sp(p, before=12, after=3)
    return p


def add_p(doc, text, bold=False, color=None, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    sp(p, before=1, after=1)
    return p


def add_divider(doc):
    p = doc.add_paragraph("─" * 72)
    p.runs[0].font.size = Pt(7)
    p.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    sp(p, before=8, after=8)


def add_glossary_entry(doc, term, definition):
    p = doc.add_paragraph()
    r_term = p.add_run(term + ": ")
    r_term.bold = True
    r_term.font.size = Pt(11)
    r_term.font.color.rgb = COLOR_TERM
    r_def = p.add_run(definition)
    r_def.font.size = Pt(11)
    r_def.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    sp(p, before=3, after=3)


# ─── Datos del glosario ──────────────────────────────────────────────────────

GLOSSARY = [
    # ── STACK Y ARQUITECTURA ──────────────────────────────────────────────
    (
        "PERN",
        "Acrónimo del stack del proyecto: PostgreSQL (base de datos) + Express (servidor) "
        "+ React (interfaz) + Node.js (entorno de ejecución). "
        "Es la variante relacional del popular MERN (que usa MongoDB en vez de PostgreSQL).",
    ),
    (
        "Stack tecnológico",
        "Conjunto de tecnologías que se usan juntas para construir la aplicación: "
        "lenguajes, frameworks, librerías y herramientas. En MyPlaythrough el stack es PERN.",
    ),
    (
        "SPA (Single Page Application)",
        "Aplicación web que carga una sola página HTML y actualiza el contenido sin recargar "
        "el navegador completo. React con React Router DOM gestiona las rutas internamente "
        "(/, /community, /settings…). El servidor solo sirve el index.html una vez.",
    ),
    (
        "API REST",
        "Interfaz de comunicación entre el cliente y el servidor usando el protocolo HTTP. "
        "Cada recurso tiene una URL (/api/games, /api/auth/me) y se operan con métodos "
        "HTTP estándar: GET (leer), POST (crear), PUT/PATCH (actualizar), DELETE (borrar). "
        "Las respuestas son JSON.",
    ),
    (
        "HTTP / métodos HTTP",
        "Protocolo de transferencia de la web. Los métodos principales son: "
        "GET (pedir datos), POST (crear), PATCH (actualizar parcialmente), "
        "PUT (actualizar completo), DELETE (borrar). "
        "La API de MyPlaythrough usa todos ellos según la operación.",
    ),
    (
        "JSON",
        "JavaScript Object Notation. Formato de texto para enviar datos entre cliente y servidor. "
        "Ejemplo: {\"nombre_usuario\": \"Tizza\", \"rol\": \"admin\"}. "
        "Toda la API de MyPlaythrough envía y recibe JSON.",
    ),
    (
        "Middleware",
        "Función que se ejecuta 'en medio' de una petición HTTP, antes de que llegue a la lógica "
        "de la ruta. En Express se encadenan: authMiddleware comprueba el JWT, "
        "adminMiddleware comprueba el rol, express.json() parsea el cuerpo. "
        "Si uno falla, devuelve un error sin pasar al siguiente.",
    ),
    (
        "Docker / Docker Compose",
        "Plataforma para empaquetar aplicaciones en contenedores. "
        "docker compose up --build levanta PostgreSQL y el servidor Express juntos, "
        "sin necesidad de instalarlos manualmente en el ordenador. "
        "Útil para que el corrector pueda probar la app con un solo comando.",
    ),
    (
        "npm",
        "Node Package Manager. Gestor de paquetes de Node.js. "
        "Instala las dependencias declaradas en package.json (express, react, tailwindcss…). "
        "npm run dev arranca el servidor en desarrollo.",
    ),
    # ── FRONTEND ─────────────────────────────────────────────────────────
    (
        "React",
        "Librería de JavaScript para construir interfaces de usuario declarativas. "
        "La UI se describe como componentes que se re-renderizan cuando cambia su estado o sus props. "
        "Versión 18 en MyPlaythrough. Los archivos tienen extensión .jsx.",
    ),
    (
        "Componente (React)",
        "Pieza reutilizable de interfaz: una función de JavaScript que devuelve JSX. "
        "Ejemplos en MyPlaythrough: GameListRow (una fila de juego), "
        "UserAvatar (el robot del usuario), NotificationBell (la campana del header). "
        "Cada componente tiene su propio estado y recibe datos por props.",
    ),
    (
        "JSX",
        "Extensión de sintaxis de JavaScript que permite escribir etiquetas HTML dentro de JS. "
        "Ejemplo: <button className=\"btn\" onClick={handleClick}>Guardar</button>. "
        "El compilador (Vite/Babel) lo convierte a llamadas React.createElement().",
    ),
    (
        "Hook (React)",
        "Funciones especiales de React que permiten usar estado y ciclo de vida en componentes "
        "funcionales (sin clases). Los más usados en MyPlaythrough: "
        "useState (gestionar un dato local), useEffect (ejecutar código al montar/actualizar), "
        "useCallback (memorizar una función), useMemo (memorizar un cálculo).",
    ),
    (
        "useState",
        "Hook de React que declara una variable de estado local en un componente. "
        "Ejemplo: const [loading, setLoading] = useState(false). "
        "Cuando llamas a setLoading(true), React re-renderiza el componente.",
    ),
    (
        "useEffect",
        "Hook de React para ejecutar código con efectos secundarios: peticiones a la API, "
        "suscripciones, timers. Se ejecuta después de renderizar y cuando cambian las dependencias "
        "del array que se le pasa.",
    ),
    (
        "Props",
        "Datos que un componente padre pasa a un componente hijo. Son de solo lectura: "
        "el hijo no puede modificarlos directamente. "
        "Ejemplo: <UserAvatar avatarId=\"robot-3\" size=\"xl\" />.",
    ),
    (
        "Vite",
        "Herramienta de build y servidor de desarrollo para proyectos web. "
        "Mucho más rápido que Create React App porque en desarrollo sirve los módulos ES nativos "
        "sin empaquetar, y en producción usa Rollup. Puerto por defecto: 5173.",
    ),
    (
        "Create React App (CRA)",
        "Herramienta anterior de Facebook/Meta para iniciar proyectos React. "
        "Más lenta en arranque y build que Vite. Se menciona como comparativa: "
        "Vite lo sustituyó como estándar porque es significativamente más rápido.",
    ),
    (
        "Tailwind CSS",
        "Framework de CSS 'utility-first': en vez de escribir clases CSS propias, "
        "combinas clases predefinidas directamente en el HTML/JSX. "
        "Ejemplo: className=\"rounded-xl border border-white/10 bg-slate-900 p-4 text-sm\". "
        "Evita tener que saltar entre archivos CSS y componentes.",
    ),
    (
        "React Router DOM",
        "Librería para gestionar la navegación dentro de una SPA. "
        "Define rutas como /community, /settings, /user/:id sin recargar la página. "
        "En MyPlaythrough: <Routes> y <Route> en App.jsx, "
        "useNavigate() para navegar por código, location.state para pasar datos entre rutas.",
    ),
    (
        "location.state (React Router)",
        "Mecanismo de React Router para pasar datos de una ruta a otra sin que aparezcan en la URL. "
        "MyPlaythrough lo usa para mostrar el aviso de confirmación en la colección "
        "al volver desde el formulario de guardar: la página de colección recibe "
        "location.state.flashGameSaved y muestra el mensaje.",
    ),
    (
        "localStorage",
        "Almacenamiento del navegador persistente entre recargas de página. "
        "MyPlaythrough guarda el JWT (token) y el objeto usuario ahí. "
        "Al cerrar sesión se borran. Diferente de sessionStorage (que dura solo la sesión del tab).",
    ),
    (
        "Skeleton loading",
        "Placeholders animados con forma de contenido que se muestran mientras los datos cargan. "
        "Mejoran la percepción de velocidad: el usuario ve que algo va a aparecer "
        "en vez de una página en blanco.",
    ),
    (
        "Tour guiado (react-joyride)",
        "Librería que muestra una secuencia de tooltips sobre elementos de la UI. "
        "En MyPlaythrough se muestra una vez tras el primer registro del usuario en ese navegador "
        "(guardado en localStorage). Se puede relanzar desde Perfil.",
    ),
    (
        "Web Audio API",
        "API del navegador para generar y reproducir audio por código (sin archivos de audio). "
        "MyPlaythrough la usa para el tono breve que suena al recibir una recomendación nueva. "
        "Los navegadores requieren que el usuario haya interactuado con la página antes de "
        "reproducir audio (restricción anti-autoplay).",
    ),
    (
        "Framer Motion",
        "Librería de animaciones para React. MyPlaythrough la usa en las diapositivas de la "
        "presentación: fadeUp (entrar con desenfoque y movimiento hacia arriba), "
        "stagger (elementos que aparecen en cascada), spring (animación elástica). "
        "Se usa solo en la presentación, no en la app principal.",
    ),
    (
        "TypeScript",
        "Superset de JavaScript que añade tipado estático. "
        "Detecta errores antes de ejecutar el código. "
        "En MyPlaythrough se usa en la mini app de presentación (.tsx). "
        "El cliente de la app principal está en JavaScript (.jsx).",
    ),
    (
        "CRUD",
        "Acrónimo de las cuatro operaciones básicas sobre datos: "
        "Create (crear), Read (leer), Update (actualizar), Delete (borrar). "
        "La colección de MyPlaythrough implementa CRUD completo sobre juegos.",
    ),
    (
        "PrintDeck",
        "Componente estático de React que replica las slides sin Framer Motion. "
        "Se usa al imprimir o guardar como PDF desde el navegador. "
        "Si se imprime el árbol animado de Framer, el PDF sale corrupto o con páginas en blanco; "
        "PrintDeck lo evita renderizando HTML plano.",
    ),
    # ── BACKEND Y NODE ────────────────────────────────────────────────────
    (
        "Node.js",
        "Entorno de ejecución de JavaScript fuera del navegador, basado en el motor V8 de Chrome. "
        "Permite usar JS en el servidor. El servidor Express de MyPlaythrough corre sobre Node.js. "
        "Versión requerida: 18+.",
    ),
    (
        "Express",
        "Framework web minimalista para Node.js. Define rutas (router.get, router.post…), "
        "middlewares y gestiona peticiones y respuestas HTTP. "
        "MyPlaythrough usa Express 5 (la versión más reciente).",
    ),
    (
        "Router (Express)",
        "Objeto que agrupa rutas relacionadas. MyPlaythrough tiene un router por tema: "
        "auth.routes.js, games.routes.js, social.routes.js, admin.routes.js… "
        "Todos se montan en index.js bajo /api/.",
    ),
    (
        "Puerto (port)",
        "Número que identifica un proceso de red en un ordenador. "
        "El servidor Express escucha en el puerto 3000, "
        "Vite en el 5173, PostgreSQL en el 5432 (nativo Windows) o 5433 (Docker).",
    ),
    (
        "Proxy de imágenes",
        "El servidor de MyPlaythrough descarga la imagen de portada desde Steam o RAWG "
        "y la reenvía al navegador del cliente. Evita errores CORS y bloqueos de hotlinking "
        "que ocurrirían si el navegador la pidiera directamente al CDN externo.",
    ),
    (
        "Hotlinking",
        "Pedir un recurso (imagen, video) directamente desde el servidor donde está alojado, "
        "sin servir desde el propio servidor. Muchos CDN lo bloquean para ahorrar ancho de banda. "
        "El proxy de MyPlaythrough evita este problema.",
    ),
    (
        "Seed / script de seed",
        "Script que inserta datos de ejemplo en la base de datos. "
        "MyPlaythrough tiene dos: seed:demo (cuenta demo con juegos reales) y "
        "seed:presentation (7 usuarios con colecciones, seguimientos y recomendaciones). "
        "Contraseña de todos: Presentacion2026!.",
    ),
    # ── AUTENTICACIÓN Y SEGURIDAD ─────────────────────────────────────────
    (
        "JWT (JSON Web Token)",
        "Estándar (RFC 7519) para crear tokens de autenticación. "
        "Al hacer login, el servidor firma un token con JWT_SECRET que contiene id, email y rol. "
        "El cliente lo guarda en localStorage y lo manda en cada petición como "
        "'Authorization: Bearer <token>'. El servidor lo verifica sin necesitar BD. "
        "Expira en 7 días. No hay estado de sesión en el servidor.",
    ),
    (
        "bcrypt / bcryptjs",
        "Algoritmo para hashear contraseñas de forma segura. "
        "Coste 10: hace el proceso deliberadamente lento (≈100ms) para dificultar ataques "
        "de fuerza bruta. La función bcrypt.compare() verifica si una contraseña corresponde "
        "al hash sin desencriptarla. Nunca se guarda la contraseña en claro.",
    ),
    (
        "Hash (de contraseña)",
        "Resultado de aplicar bcrypt a la contraseña. Es una cadena opaca e irreversible. "
        "Para verificar si una contraseña es correcta, se aplica bcrypt.compare(contraseña, hash): "
        "si coincide devuelve true. El atacante que robe la BD no puede recuperar la contraseña.",
    ),
    (
        "authMiddleware",
        "Middleware de Express que comprueba que la petición lleva un JWT válido "
        "en la cabecera Authorization: Bearer <token>. "
        "Si es válido, adjunta el payload (id, email, rol) en req.user y continúa. "
        "Si falta o es inválido → HTTP 401.",
    ),
    (
        "adminMiddleware",
        "Middleware que, tras authMiddleware, consulta la BD para comprobar que el usuario "
        "tiene rol 'admin'. Lo comprueba en BD (no solo en el token) por si el rol cambió "
        "después de emitir el token. Si no es admin → HTTP 403.",
    ),
    (
        "Rol (user / admin)",
        "Nivel de permisos de un usuario. 'user' es el rol por defecto de todos los registrados. "
        "'admin' puede gestionar usuarios, fichas y LFG desde el panel. "
        "El rol se guarda en la columna 'rol' de la tabla usuarios y se comprueba en BD "
        "en cada petición al panel de administración.",
    ),
    (
        "Rate limiting (límite de frecuencia)",
        "Técnica para limitar el número de peticiones que una misma IP puede hacer en un tiempo. "
        "En MyPlaythrough: express-rate-limit en /api/auth/login y /api/auth/register. "
        "Máximo 40 intentos en 15 minutos. Si se supera → HTTP 429 (Too Many Requests). "
        "Evita ataques de fuerza bruta en el login.",
    ),
    (
        "CORS (Cross-Origin Resource Sharing)",
        "Mecanismo del navegador que bloquea peticiones JavaScript a un dominio distinto al origen. "
        "MyPlaythrough configura el middleware cors() con CORS_ORIGIN (variable de entorno) "
        "para que solo el frontend (localhost:5173 en desarrollo) pueda llamar a la API. "
        "Sin esto, cualquier web podría llamar a la API desde el navegador del usuario.",
    ),
    (
        "Consultas parametrizadas",
        "Forma de escribir SQL donde los valores del usuario van como parámetros ($1, $2) "
        "y nunca se concatenan en el texto de la query. "
        "Ejemplo: SELECT * FROM usuarios WHERE email = $1 con [email] como parámetros. "
        "Evita la inyección SQL porque el motor SQL trata los parámetros como datos, "
        "nunca como código.",
    ),
    (
        "Inyección SQL (SQL injection)",
        "Ataque que consiste en introducir código SQL en un campo de formulario para que el "
        "servidor lo ejecute y manipule la base de datos (borrar tablas, robar datos…). "
        "MyPlaythrough lo previene usando consultas parametrizadas en todas las queries.",
    ),
    (
        "HTTP 401 / 403 / 429",
        "Códigos de estado HTTP de error: "
        "401 = No autenticado (falta token o es inválido). "
        "403 = Prohibido (estás autenticado pero no tienes permiso, ej: no eres admin). "
        "429 = Demasiadas peticiones (rate limit alcanzado). "
        "404 = No encontrado. 500 = Error interno del servidor.",
    ),
    (
        "Variables de entorno (.env)",
        "Archivo de configuración que guarda valores sensibles (contraseñas, claves de API, secretos) "
        "fuera del código. En MyPlaythrough: server/.env contiene DB_*, JWT_SECRET, "
        "CORS_ORIGIN y RAWG_API_KEY. No se sube al repositorio git.",
    ),
    # ── BASE DE DATOS ─────────────────────────────────────────────────────
    (
        "PostgreSQL",
        "Sistema de gestión de base de datos relacional de código abierto. "
        "Versión 16 en MyPlaythrough. Soporta: claves foráneas, CASCADE delete, "
        "índices únicos con expresión, transacciones ACID, y el tipo SERIAL para auto-incremento. "
        "Escucha en el puerto 5432 (instalación nativa) o 5433 (Docker Compose).",
    ),
    (
        "Base de datos relacional",
        "Organiza los datos en tablas con filas y columnas, y define relaciones entre tablas "
        "mediante claves foráneas. Garantiza integridad referencial: no puedes tener "
        "una fila en 'juegos' que apunte a un usuario que no existe.",
    ),
    (
        "Clave primaria (PK)",
        "Columna o combinación de columnas que identifica de forma única cada fila de una tabla. "
        "En MyPlaythrough la mayoría de tablas tienen id SERIAL PRIMARY KEY (auto-incremental). "
        "juego_comentario_votos y usuario_seguimientos tienen PK compuesta (dos columnas).",
    ),
    (
        "Clave foránea (FK)",
        "Columna que referencia la clave primaria de otra tabla. "
        "Ejemplo: juegos.usuario_id referencia usuarios.id. "
        "Garantiza integridad referencial: no puedes insertar un juego para un usuario que no existe.",
    ),
    (
        "CASCADE delete",
        "Comportamiento de una FK: cuando se borra la fila referenciada, "
        "se borran automáticamente todas las filas que la referencian. "
        "En MyPlaythrough: borrar un usuario → se borran sus juegos, comentarios, "
        "seguimientos y recomendaciones automáticamente. Evita datos huérfanos.",
    ),
    (
        "Normalización (base de datos)",
        "Proceso de diseño que organiza las tablas para evitar duplicar datos. "
        "En MyPlaythrough: la tabla catalogo_juegos existe para guardar título y portada "
        "una sola vez, aunque 100 usuarios añadan el mismo juego. "
        "Cada uno tiene su fila en 'juegos' con sus datos personales, "
        "pero comparten la fila de catalogo_juegos.",
    ),
    (
        "Índice único con expresión",
        "Índice de PostgreSQL que garantiza unicidad aplicando una función al valor. "
        "En MyPlaythrough: CREATE UNIQUE INDEX idx_usuarios_nombre_lower "
        "ON usuarios (LOWER(TRIM(nombre_usuario))). "
        "Hace que 'Tizza', 'tizza' y 'TIZZA' sean el mismo nombre a efectos de unicidad.",
    ),
    (
        "SERIAL",
        "Tipo de PostgreSQL que auto-incrementa el valor de una columna con cada nueva fila. "
        "Equivale a un INTEGER con una secuencia automática. "
        "Se usa como id en la mayoría de tablas de MyPlaythrough.",
    ),
    (
        "Tabla usuarios",
        "Tabla principal de MyPlaythrough. Columnas relevantes: id, nombre_usuario (único), "
        "email (único), password_hash (bcrypt), rol (user/admin), "
        "avatar_id (qué robot SVG tiene), notificaciones_sonido, fecha_registro.",
    ),
    (
        "Tabla juegos",
        "Biblioteca personal de un usuario. Una fila por juego que ha añadido. "
        "Columnas: id, usuario_id (FK→usuarios), titulo, estado "
        "(jugando/completado/pendiente/abandonado/wishlist), plataforma, "
        "puntuacion, horas_jugadas, comentario, url_imagen, catalogo_id (FK→catalogo_juegos).",
    ),
    (
        "Tabla catalogo_juegos",
        "Ficha compartida de un juego. Columnas: id, titulo, url_imagen, "
        "rawg_id (UNIQUE), steam_app_id (UNIQUE). "
        "Varios usuarios pueden tener filas en 'juegos' apuntando a la misma fila de catalogo_juegos. "
        "El título y la portada se guardan una sola vez.",
    ),
    (
        "Tabla usuario_seguimientos",
        "Implementa el grafo social de 'seguir'. "
        "PK compuesta (seguidor_id, seguido_id), ambas FK a usuarios. "
        "Una fila = A sigue a B.",
    ),
    (
        "Tabla juego_recomendaciones",
        "Mensajes de recomendación de juego entre usuarios. "
        "remitente_id y destinatario_id (FK→usuarios), juego_id (FK→juegos), "
        "mensaje, leida (booleano), fecha.",
    ),
    (
        "Tabla lfg_publicaciones",
        "Publicaciones de 'buscar grupo' (LFG). "
        "usuario_id (FK→usuarios), juego_id (FK→juegos), "
        "modo (online/local/otro), descripcion, activa, fecha.",
    ),
    (
        "Tabla juego_comentarios",
        "Comentarios sobre un juego, en formato de hilo. "
        "parent_id es FK a la misma tabla (juego_comentarios) para implementar respuestas. "
        "Si parent_id es NULL, es un comentario raíz.",
    ),
    (
        "Tabla juego_comentario_votos",
        "Votos sobre comentarios de primer nivel. "
        "PK compuesta (comentario_id, usuario_id). "
        "Columna valor: SMALLINT con CHECK (valor IN (-1, 1)). "
        "Un usuario solo puede votar una vez por comentario.",
    ),
    # ── APIS EXTERNAS ─────────────────────────────────────────────────────
    (
        "Steam API",
        "API pública de Valve para buscar información de videojuegos (títulos, portadas, IDs). "
        "MyPlaythrough la usa en el buscador de portadas. "
        "No requiere clave API para la búsqueda básica de carátulas.",
    ),
    (
        "RAWG API",
        "API pública de base de datos de videojuegos (rawg.io). "
        "Clave gratuita con 20.000 peticiones/mes. Opcional pero recomendada: "
        "sin ella el buscador solo encuentra juegos de Steam, "
        "con RAWG también aparecen juegos de Nintendo, PlayStation, etc. "
        "Se configura en server/.env como RAWG_API_KEY.",
    ),
    # ── FUNCIONES Y UX ────────────────────────────────────────────────────
    (
        "LFG (Looking for Group)",
        "Funcionalidad para publicar que buscas con quién jugar a un juego. "
        "La publicación está vinculada a un juego de tu biblioteca "
        "e incluye el modo (online, co-op local, otro) y una descripción. "
        "Visible en la sección Comunidad.",
    ),
    (
        "Accesibilidad (a11y)",
        "Conjunto de prácticas para que la app sea usable por personas con diversidad funcional. "
        "En MyPlaythrough: skip link (enlace para saltar al contenido principal), "
        "foco automático al elemento <main> al cambiar de vista, "
        "y labels en todos los controles de vista y ordenación.",
    ),
    (
        "Panel de administración",
        "Sección accesible solo para usuarios con rol 'admin' en /admin. "
        "Permite gestionar usuarios (ver lista, borrar cuenta con confirmación reforzada), "
        "fichas de juego de cualquier usuario, y publicaciones LFG. "
        "Si un usuario sin permisos intenta acceder a /admin, la app le redirige al inicio.",
    ),
    (
        "Modal de confirmación reforzado",
        "Diálogo que exige al admin escribir el nombre exacto del usuario antes de borrarlo. "
        "Previene borrados accidentales con un solo clic. "
        "También aparece en el modal de enviar recomendación: "
        "muestra el nombre del destinatario y el título del juego antes de confirmar.",
    ),
    (
        "Avatar (robot SVG)",
        "Los 10 avatares predefinidos de MyPlaythrough son ilustraciones SVG de robots. "
        "Se guardan en la columna avatar_id de la tabla usuarios (ej: 'robot-0', 'robot-5'). "
        "El usuario elige el suyo en /settings y se muestra en el header, "
        "perfiles públicos, comunidad y comentarios.",
    ),
]


# ─── Main ─────────────────────────────────────────────────────────────────────

def main():
    if not DOCX.exists():
        print(f"ERROR: no se encuentra {DOCX}")
        return

    print(f"Abriendo: {DOCX}")
    doc = Document(str(DOCX))

    print("1. Corrigiendo colores claros...")
    fix_light_colors(doc)

    # Comprobar si el glosario ya existe para no duplicarlo
    already_has_glossary = any(
        "GLOSARIO" in (p.text or "")
        for p in doc.paragraphs
    )
    if already_has_glossary:
        print("2. Glosario ya presente — no se duplica.")
        try:
            doc.save(str(DOCX))
            print(f"\nOK Guardado (solo colores): {DOCX}")
        except PermissionError:
            alt = DOCX.with_name("guion-defensa-8min-actualizado.docx")
            doc.save(str(alt))
            print(f"\nArchivo abierto en Word! Guardado como: {alt}")
        return

    print("2. Añadiendo glosario...")
    add_divider(doc)
    add_h(doc, "GLOSARIO — Conceptos clave del proyecto", 1)
    add_p(
        doc,
        "Definiciones en el contexto de MyPlaythrough. "
        "Para cada término: qué es, por qué se usa en este proyecto y cómo aparece en el guion.",
        color=RGBColor(0x4B, 0x55, 0x63),
        size=10,
    )

    sections = [
        ("Stack y arquitectura",
         ["PERN", "Stack tecnológico", "SPA (Single Page Application)",
          "API REST", "HTTP / métodos HTTP", "JSON",
          "Middleware", "Docker / Docker Compose", "npm"]),
        ("Frontend",
         ["React", "Componente (React)", "JSX", "Hook (React)", "useState", "useEffect",
          "Props", "Vite", "Create React App (CRA)", "Tailwind CSS",
          "React Router DOM", "location.state (React Router)", "localStorage",
          "Skeleton loading", "Tour guiado (react-joyride)",
          "Web Audio API", "Framer Motion", "TypeScript", "CRUD", "PrintDeck"]),
        ("Backend y Node",
         ["Node.js", "Express", "Router (Express)", "Puerto (port)",
          "Proxy de imágenes", "Hotlinking", "Seed / script de seed"]),
        ("Autenticación y seguridad",
         ["JWT (JSON Web Token)", "bcrypt / bcryptjs", "Hash (de contraseña)",
          "authMiddleware", "adminMiddleware", "Rol (user / admin)",
          "Rate limiting (límite de frecuencia)", "CORS (Cross-Origin Resource Sharing)",
          "Consultas parametrizadas", "Inyección SQL (SQL injection)",
          "HTTP 401 / 403 / 429", "Variables de entorno (.env)"]),
        ("Base de datos",
         ["PostgreSQL", "Base de datos relacional",
          "Clave primaria (PK)", "Clave foránea (FK)", "CASCADE delete",
          "Normalización (base de datos)", "Índice único con expresión", "SERIAL",
          "Tabla usuarios", "Tabla juegos", "Tabla catalogo_juegos",
          "Tabla usuario_seguimientos", "Tabla juego_recomendaciones",
          "Tabla lfg_publicaciones", "Tabla juego_comentarios",
          "Tabla juego_comentario_votos"]),
        ("APIs externas",
         ["Steam API", "RAWG API"]),
        ("Funciones y UX",
         ["LFG (Looking for Group)", "Accesibilidad (a11y)",
          "Panel de administración", "Modal de confirmación reforzado",
          "Avatar (robot SVG)"]),
    ]

    glossary_dict = {term: definition for term, definition in GLOSSARY}

    for section_title, terms in sections:
        add_h(doc, section_title, 2)
        for term in terms:
            if term in glossary_dict:
                add_glossary_entry(doc, term, glossary_dict[term])
            else:
                print(f"  AVISO: término sin definición: {term}")

    # Guardar
    try:
        doc.save(str(DOCX))
        print(f"\nOK Guardado: {DOCX}")
    except PermissionError:
        alt = DOCX.with_name("guion-defensa-8min-actualizado.docx")
        doc.save(str(alt))
        print(f"\nArchivo abierto en Word! Guardado como: {alt}")


if __name__ == "__main__":
    main()
